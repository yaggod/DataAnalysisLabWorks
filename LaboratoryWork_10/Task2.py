import docx
from argparse import ArgumentParser
parser = ArgumentParser()

parser.add_argument("--docx", default="result.docx", help="File to modify")
parser.add_argument("--image", required=True, help="Image file")
parser.add_argument("--text", default="", help="Text under the image")
args = parser.parse_args() 

filename = args.docx
document = docx.Document(filename)
document.add_picture(args.image)
document.add_paragraph(args.text)

document.save(filename)