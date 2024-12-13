import docx

document = docx.Document()

document.add_paragraph("В микроконтроллерах ATMega, используемых на платформах Arduino, существует три вида памяти:")
document.add_paragraph("Флеш-память: используется для хранения скетчей", "List Bullet")

listItem = document.add_paragraph("ОЗУ (", "List Bullet")
listItem.add_run("SRAM").bold = True
listItem.add_run(" - static random access memory").italic = True
listItem.add_run(", статическая оперативная память с произвольным доступом): используется для хранения и работы переменных")

document.add_paragraph("EEPREOM (энергонезависимая память): используется для хранения постоянной инфромации", "List Bullet")

document.add_paragraph('Флеш-память и EEPROM являются энергонезависимыми'
                       'видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.')


table = document.add_table(4, 5, "Table Grid")

data = [ ["", "ATmega168", "ATmega328", "ATmega1280", "ATmega2560"],
        ["Flash (1 кБ flash-памяти занят загрузчиком)", "16 Кбайт", "32 Кбайт", "128 Кбайт", "256 Кбайт"],
        ["SRAM", "1 Кбайт", "2 Кбайт", "8 Кбайт", "8 Кбайт"],
        ["EEPROM", "512 байт", "1024 байта", "4 Кбайт", "4 Кбайт"]
]

for rowIndex in range(len(data)):
    row = data[rowIndex]
    for columnIndex in range(len(row)):
        value = row[columnIndex]
        cell = table.cell(rowIndex, columnIndex)
        cell.text = value
        if (rowIndex == 0 or columnIndex == 0): # rowIndex * columnIndex == 0 but that's less readible
           cell.paragraphs[0].runs[0].bold = True

text = ('Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным'
'циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25 С.'
'Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных'
'не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально'
'щадящими по отношению к EEPROM.')

document.add_paragraph(text).runs[0].italic = True

document.save("result.docx")

